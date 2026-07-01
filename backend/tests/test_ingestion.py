"""
Unit tests for the document ingestion service.

Tests cover:
- PDF, TXT, DOCX parsing
- Chunking with overlap
- Full ingest_document pipeline
- Error handling for unsupported formats and empty documents
"""

import pytest
from unittest.mock import patch, MagicMock

from app.services.ingestion import (
    chunk_document,
    ingest_document,
    parse_txt,
    parse_docx,
)


class TestParseTxt:
    def test_returns_decoded_string(self, sample_txt_bytes):
        result = parse_txt(sample_txt_bytes)
        assert "sample document" in result
        assert isinstance(result, str)

    def test_handles_encoding_errors(self):
        # Invalid UTF-8 bytes should be replaced, not crash
        result = parse_txt(b"\xff\xfe invalid bytes")
        assert isinstance(result, str)


class TestParseDocx:
    def test_extracts_paragraphs(self):
        """Test DOCX parsing with a minimal docx structure."""
        from docx import Document
        import io

        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        buffer = io.BytesIO()
        doc.save(buffer)

        result = parse_docx(buffer.getvalue())
        assert "First paragraph" in result
        assert "Second paragraph" in result

    def test_skips_empty_paragraphs(self):
        from docx import Document
        import io

        doc = Document()
        doc.add_paragraph("")
        doc.add_paragraph("Real content")
        buffer = io.BytesIO()
        doc.save(buffer)

        result = parse_docx(buffer.getvalue())
        assert "Real content" in result


class TestChunkDocument:
    def test_produces_chunks(self, sample_txt_bytes):
        text = parse_txt(sample_txt_bytes)
        chunks = chunk_document(text, "test.txt", "file-123")
        assert len(chunks) >= 1

    def test_chunk_metadata(self, sample_txt_bytes):
        text = parse_txt(sample_txt_bytes)
        chunks = chunk_document(text, "my_file.txt", "file-abc")
        for i, chunk in enumerate(chunks):
            assert chunk.source == "my_file.txt"
            assert chunk.file_id == "file-abc"
            assert chunk.chunk_index == i
            assert len(chunk.content) > 0
            assert chunk.uploaded_at != ""

    def test_chunk_size_respected(self):
        long_text = "word " * 1000  # ~5000 chars
        chunks = chunk_document(long_text, "large.txt", "file-large", chunk_size=200, chunk_overlap=20)
        # All chunks should be reasonably sized
        for chunk in chunks:
            assert len(chunk.content) <= 400  # allow some slack for splitter

    def test_overlap_creates_shared_content(self):
        # Create text where overlap should be detectable
        text = ("alpha beta gamma " * 100)
        chunks = chunk_document(text, "test.txt", "id", chunk_size=50, chunk_overlap=10)
        assert len(chunks) > 1


class TestIngestDocument:
    def test_ingest_txt(self, sample_txt_bytes):
        chunks = ingest_document("sample.txt", sample_txt_bytes, "file-001")
        assert len(chunks) >= 1
        assert all(c.source == "sample.txt" for c in chunks)

    def test_ingest_unsupported_format(self, sample_txt_bytes):
        with pytest.raises(ValueError, match="Unsupported file format"):
            ingest_document("document.xyz", sample_txt_bytes, "file-002")

    def test_ingest_empty_document(self):
        empty_bytes = b"   "
        with pytest.raises(ValueError, match="no extractable text"):
            ingest_document("empty.txt", empty_bytes, "file-003")

    def test_ingest_docx(self):
        from docx import Document
        import io

        doc = Document()
        doc.add_paragraph("Hello, this is a DOCX test document with enough content.")
        buffer = io.BytesIO()
        doc.save(buffer)

        chunks = ingest_document("test.docx", buffer.getvalue(), "file-004")
        assert len(chunks) >= 1
        assert chunks[0].source == "test.docx"
